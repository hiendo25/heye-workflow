# -*- coding: utf-8 -*-
"""Phần 1: bìa, mục lục, kết luận sớm, và nhóm biểu đồ NHÂN SỰ."""

from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build(g):
    doc, h, p, table, img, note, pagebreak, field = (
        g.doc, g.h, g.p, g.table, g.img, g.note, g.pagebreak, g.field
    )
    BRAND, MUTED = g.BRAND, g.MUTED

    # ═══════════ BÌA ═══════════
    for _ in range(5):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Biểu đồ thống kê công việc và nhân sự")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = BRAND

    s_ = doc.add_paragraph()
    s_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s_.add_run("Nghiên cứu 8 sản phẩm — tham chiếu cho màn Thống kê của HeyE")
    r.font.size = Pt(13)
    r.font.color.rgb = MUTED

    for _ in range(6):
        doc.add_paragraph()

    tb = doc.add_table(rows=4, cols=2)
    tb.alignment = 1
    for i, (k, v) in enumerate([
        ("Phạm vi", "Jira · Asana · ClickUp · Monday · Linear · Productive · Harvest · Float"),
        ("Bằng chứng", "39 ảnh thật tải từ tài liệu chính thức, đã xem trực tiếp"),
        ("Ngày", "24/08/2026"),
        ("Đối tượng đọc", "PM, BA, người quyết định làm màn Thống kê"),
    ]):
        c0, c1 = tb.rows[i].cells
        c0.text = ""
        rr = c0.paragraphs[0].add_run(k)
        rr.bold = True
        rr.font.size = Pt(10)
        c1.text = ""
        rr = c1.paragraphs[0].add_run(v)
        rr.font.size = Pt(10)
        c0.width, c1.width = Cm(3.6), Cm(12.4)

    # ═══════════ KẾT LUẬN SỚM ═══════════
    pagebreak()
    h("Kết luận trước — 5 biểu đồ HeyE nên làm", 1)
    p(
        "Đặt ngay đầu tài liệu để người bận đọc xong phần này là đủ quyết định. "
        "Phần sau giải thích vì sao chọn như vậy."
    )
    table(
        ["#", "Biểu đồ", "Trả lời câu hỏi", "Học từ", "Dữ liệu HeyE"],
        [
            [
                "1",
                "**Lưới tải nhân sự**\nWorkload grid",
                "Tuần sau ai quá tải, ai rảnh để nhận thêm việc",
                "Monday · Asana",
                "✓ có đủ",
            ],
            [
                "2",
                "**Tỷ lệ giờ ra tiền**\nBillable utilization",
                "Giờ làm có bán được không, ai đang chạy không tải",
                "Productive.io",
                "✓ có đủ",
            ],
            [
                "3",
                "**Hàng số to**\nKPI tile",
                "Tổng quan một dòng cho sếp: việc, giờ, tiền",
                "ClickUp",
                "✓ có đủ",
            ],
            [
                "4",
                "**Tạo mới vs Hoàn thành**\nCreated vs Resolved",
                "Đội đang đuổi kịp hay đuối so với việc đổ vào",
                "Jira",
                "✓ có đủ",
            ],
            [
                "5",
                "**Sản lượng theo kỳ**\nVelocity",
                "Tháng này so tháng trước, đang tốt lên hay xấu đi",
                "Jira · ClickUp",
                "✓ có đủ",
            ],
        ],
        widths=[0.9, 3.6, 5.4, 2.6, 2.1],
        small=True,
    )
    note(
        "Cả 5 đều tính được từ dữ liệu HeyE đang có, không cần thêm bảng mới. "
        "Riêng biểu đồ 1 và 2 cần ngày lễ Việt Nam, hiện đang gắn cứng trong mã "
        "nguồn — nên tách ra thành cấu hình.",
    )

    h("Ba điều đáng chú ý nhất", 2)
    field(
        "Không ai đo năng suất từng cá nhân",
        "cả Jira lẫn Asana đều KHÔNG có biểu đồ năng suất cá nhân theo thời gian. "
        "Velocity của Jira là của cả đội. Workload của Asana đo tải sắp tới, không "
        "đo thành tích đã qua. Đây gần như chắc chắn là lựa chọn có chủ đích: xếp "
        "hạng nhân viên bằng số task đếm được sẽ khiến người ta chẻ nhỏ task cho "
        "đẹp số.",
    )
    field(
        "Không ai đặt tiền lên hàng KPI chính",
        "trong 8 sản phẩm khảo sát, không sản phẩm nào đưa chi phí hay lợi nhuận "
        "lên hàng số to. HeyE có sẵn dữ liệu này — đây là chỗ khác biệt được.",
    )
    field(
        "Rất ít đầu tư cho xu hướng dài hạn",
        "chỉ có 3 biểu đồ thật sự đa kỳ trong toàn khảo sát. Phần lớn widget nói "
        "về sprint hiện tại. Làm tốt xu hướng nhiều tháng cũng là chỗ khác biệt được.",
    )

    # ═══════════ MỤC LỤC ═══════════
    pagebreak()
    h("Mục lục", 1)
    table(
        ["#", "Nội dung", "Tóm tắt"],
        [
            ["1", "**Nhân sự và năng suất**", "Utilization, capacity, quá tải — phần quan trọng nhất"],
            ["2", "**Tiến độ công việc**", "Burndown, burnup, velocity, cumulative flow"],
            ["3", "**Hàng số to và thanh tiến độ**", "KPI tile, battery, số tổng"],
            ["4", "**Cách phân tầng sếp và PM**", "Đặt biểu đồ nào ở đâu"],
            ["5", "**Quy ước màu và ngưỡng cảnh báo**", "Khi nào tô đỏ"],
            ["6", "**Bẫy khi làm cho đội nhỏ**", "Biểu đồ nào vô dụng với 5–10 người"],
            ["7", "**Đối chiếu với dữ liệu HeyE**", "Tính được gì ngay, thiếu gì"],
        ],
        widths=[1.2, 6.5, 8.3],
    )

    # ═══════════ 1. NHÂN SỰ ═══════════
    pagebreak()
    h("1. Nhân sự và năng suất", 1)
    p(
        "Phần quan trọng nhất. Bốn sản phẩm cùng gọi là “utilization” nhưng định "
        "nghĩa MẪU SỐ khác nhau, ra số khác hẳn."
    )

    h("1.1 Bốn cách tính utilization — chọn cái nào", 2)
    table(
        ["Sản phẩm", "Công thức", "Mẫu số là gì", "Trừ lễ và phép?"],
        [
            [
                "**Productive.io**",
                "Billable / Available",
                "Available = Capacity − nghỉ phép",
                "**Có** — trừ cả lễ và phép",
            ],
            ["Harvest", "Billable / Total available", "Available hours", "Có"],
            [
                "Toggl Track",
                "Billable / Scheduled",
                "Giờ theo lịch cài sẵn",
                "**Không** — lịch tĩnh",
            ],
            [
                "Float",
                "Billable / Scheduled",
                "Giờ đã phân bổ",
                "Capacity có trừ, nhưng mẫu số là Scheduled",
            ],
        ],
        widths=[3.0, 3.6, 4.6, 4.8],
        small=True,
    )
    p("Trích nguyên văn tài liệu Productive:")
    note(
        "“In Productive, utilization is based on availability rather than "
        "capacity, which gives a more accurate measure.”",
    )
    p("Chuỗi ba tầng cần nhớ:")
    table(
        ["Tầng", "Công thức", "Ví dụ thật từ tài liệu Productive"],
        [
            ["Capacity", "lịch làm việc − ngày lễ", "5 ngày × 8h × 20 ngày = 160h, trừ 2 ngày lễ → **144h**"],
            ["Available", "Capacity − nghỉ phép", "176h − 5 ngày phép (40h) → **136h**"],
            ["Utilization", "Billable / Available", "80h / 112h × 100 = **71%**"],
        ],
        widths=[2.6, 4.4, 9.0],
    )
    note(
        "Khuyến nghị cho HeyE: dùng mẫu số Available kiểu Productive. Đây là mẫu "
        "số duy nhất không phạt oan người nghỉ phép — dùng lịch tĩnh kiểu Toggl "
        "thì người nghỉ 10 ngày tụt xuống 50% dù họ làm rất chăm những ngày có mặt.",
    )

    h("1.2 Cạm bẫy Float — công thức cho ra màn hình vô dụng", 2)
    img("float-people-1.png", "Hình 1 — Báo cáo People của Float")
    p(
        "Nhìn cột Billable %: mọi người đều 100%, dù Capacity 136h mà Scheduled "
        "chỉ 61.2h. Vì 100% giờ được phân bổ là billable."
    )
    note(
        "Đây là chỉ số “chất lượng công việc được giao”, KHÔNG phải “mức bận”. "
        "Nếu HeyE copy nhầm công thức này sẽ ra màn hình ai cũng 100% và không "
        "phát hiện được ai quá tải.",
        "warn",
    )

    h("1.3 Hai chỉ số utilization phải có cả hai", 2)
    p("Harvest phân biệt rõ nhất:")
    field(
        "Billable utilization",
        "phần trăm thời gian làm việc dành cho việc trực tiếp sinh doanh thu. "
        "Công thức: giờ tính tiền chia giờ khả dụng.",
    )
    field(
        "Total utilization",
        "gồm cả việc nội bộ: họp, đào tạo, hành chính. Công thức: toàn bộ giờ đã "
        "làm chia giờ khả dụng.",
    )
    p("Khoảng cách giữa hai chỉ số này chính là chẩn đoán:")
    table(
        ["Total", "Billable", "Nghĩa là", "Việc cần làm"],
        [
            ["95%", "45%", "Làm quần quật nhưng vào việc nội bộ", "Vấn đề phân công, không phải thiếu việc"],
            ["50%", "45%", "Thật sự rảnh", "Vấn đề thiếu việc"],
            ["> 100%", "—", "Quá tải thật", "Giảm tải ngay"],
        ],
        widths=[1.8, 1.8, 5.6, 6.8],
    )
    note("Một chỉ số đơn lẻ không phân biệt được hai ca đầu.")

    # ─────────
    pagebreak()
    h("1.4 Lưới tải nhân sự — mô hình đáng học nhất", 2)
    img("mo-workload-i14.png", "Hình 2 — Workload của Monday, tooltip “4 out of 3”")
    p(
        "Đây KHÔNG phải biểu đồ cột mà là lưới bong bóng: mỗi dòng một người, mỗi "
        "cột một tuần, mỗi ô một bong bóng."
    )
    field("Kích thước bong bóng", "khối lượng việc.")
    field("Màu", "xanh dương là trong ngưỡng, **đỏ là quá tải**.")
    field("Số góc dưới phải", "giá trị thực — số task hoặc số giờ.")
    field(
        "Tooltip khi hover",
        "“4 out of 3” — sức chứa 3, đang gánh 4, kèm danh sách task và effort "
        "từng cái. Cách nói này dễ hiểu hơn hẳn “133%”.",
    )
    field("Dòng Unassigned", "việc chưa giao ai, nằm cuối lưới.")
    note(
        "Đây là biểu đồ DUY NHẤT trả lời được câu “tuần sau ai rảnh”. ClickUp và "
        "Linear chỉ nói được hiện tại. Với người quản lý, biết trước mới xoay kịp.",
    )

    img("mo-workload-i13.png", "Hình 3 — Ngày lễ tô gạch chéo đỏ, cuối tuần gạch xám")
    p(
        "Monday tô gạch chéo cho ngày lễ kèm icon tam giác cảnh báo khi có task rơi "
        "vào ngày nghỉ. Rất hợp bối cảnh Việt Nam: Tết, 30/4, 2/9."
    )

    img("asana-workload-overview.png", "Hình 4 — Workload của Asana, cách vẽ khác")
    p(
        "Asana làm cùng bài toán nhưng vẽ khác: mỗi người một hàng biểu đồ vùng "
        "nhỏ, có đường ngang nét đứt đỏ đánh dấu ngưỡng, phần vượt tô đỏ. Kéo thả "
        "task ngay trên biểu đồ để giao lại người khác."
    )
    img("asana-workload-set-capacity.png", "Hình 5 — Đặt ngưỡng capacity trong Asana")
    p("Mặc định 40 giờ mỗi người mỗi tuần, cho sửa riêng từng người.")

    h("1.5 Bẫy quy đổi capacity theo khung thời gian", 2)
    p("Tài liệu của chính Monday phải cảnh báo chỗ này:")
    table(
        ["Xem theo", "Capacity được hiểu là", "Ví dụ với capacity 20h/tuần"],
        [
            ["Tuần", "đúng con số đã đặt", "20h"],
            ["Ngày", "**chia đều cho 5 ngày làm việc**", "4h mỗi ngày"],
            ["Tháng", "**nhân 4 tuần**", "80h"],
        ],
        widths=[2.6, 6.0, 7.4],
    )
    note(
        "Nếu HeyE cho đổi khung thời gian, phải nói rõ quy đổi này trên giao diện. "
        "Không thì người dùng đặt 20h/tuần rồi xem theo tháng sẽ thấy con số lạ.",
        "warn",
    )

    h("1.6 Ba cách hiện quá tải", 2)
    table(
        ["Cách", "Sản phẩm", "Đánh giá"],
        [
            [
                "**Cột giờ vượt tuyệt đối**",
                "Float",
                "**Rõ nhất.** “Vượt 3 giờ” hành động được ngay, không phải tính ngược",
            ],
            ["Bong bóng đỏ có tooltip", "Monday", "Trực quan, nhìn một giây biết ai đỏ"],
            ["Gradient màu đậm dần", "Productive", "Đẹp cho lưới lịch nhưng không đọc được con số"],
        ],
        widths=[4.2, 2.6, 9.2],
    )
    note(
        "Float tách “Còn rảnh” và “Vượt giờ” thành HAI CỘT DƯƠNG riêng, không dùng "
        "một cột có thể âm. Sạch hơn nhiều — cột âm dương lẫn lộn khó quét mắt và "
        "khó tô màu.",
    )

    h("1.7 Capacity với người bán thời gian và người vào giữa tháng", 2)
    img("prod-cap-1.png", "Hình 6 — Form cost rate của Productive")
    p(
        "Cấu trúc này gần như trùng khớp với bảng cost_rates của HeyE — lưới 7 ô "
        "Mon–Sun nhập giờ riêng, ngày bắt đầu và kết thúc, lịch nghỉ lễ."
    )
    field(
        "Bán thời gian",
        "không dùng hệ số phần trăm. Nhập thẳng giờ từng ngày, ví dụ Mon 8 · Tue 8 "
        "· Wed 4 · Thu 0 · Fri 0. Capacity tự đúng, không cần logic riêng.",
    )
    field(
        "Vào giữa tháng",
        "ngày bắt đầu cắt capacity từ ngày vào. Người vào 15/03 chỉ tính capacity "
        "nửa sau tháng 3, nên utilization tháng đầu không bị bóp méo. Đây là lỗi "
        "kinh điển nếu không xử lý.",
    )
    field(
        "Ngày lễ",
        "trừ vào Capacity. Nghỉ phép trừ tiếp ra Available.",
    )

    img("prod-cap-4.png", "Hình 7 — Biểu đồ Capacity và Availability")
    p(
        "Cột vàng Available 136h cạnh cột đỏ Capacity 176h. Khoảng hở giữa hai cột "
        "chính là thời gian nghỉ. HeyE nên copy dạng này."
    )

    h("1.8 Ngưỡng cảnh báo — số cụ thể", 2)
    p("Toggl là sản phẩm duy nhất công bố ngưỡng có màu rõ ràng:")
    table(
        ["Mức", "Màu", "Nghĩa"],
        [
            ["≥ 80%", "bình thường", "Đạt mục tiêu"],
            ["71–79%", "**vàng nhạt**", "Dưới mục tiêu"],
            ["< 70%", "**đỏ nhạt**", "Thấp hơn hẳn mục tiêu"],
        ],
        widths=[2.6, 3.4, 10.0],
    )
    note(
        "Lưu ý hướng cảnh báo: Toggl tô đỏ khi utilization THẤP, không phải cao. "
        "Họ coi utilization thấp là mất tiền.",
    )

    p("Float phân ba trạng thái, trong đó trạng thái giữa rất tinh tế:")
    field("Xám — trong ngưỡng", "dưới 100% và không có giờ vượt.")
    field(
        "Cam — phân bổ lệch",
        "tổng dưới 100% NHƯNG vẫn có giờ vượt. Nghĩa là thứ Hai cày 12h, thứ Sáu "
        "rảnh. Nhìn tổng tháng thì khỏe mạnh nhưng người đó vẫn cày đêm.",
    )
    field("Đỏ — vượt sức chứa", "trên 100%, không còn giờ trống.")
    note(
        "Trạng thái Cam là lý do phải tính theo TUẦN, không tính theo tháng. Tính "
        "theo tháng sẽ giấu mất ca phân bổ lệch.",
        "warn",
    )

    note(
        "Productive cho đặt MỤC TIÊU RIÊNG từng người, tô màu theo hiệu số giữa "
        "thực tế và mục tiêu: xanh khi vượt, đỏ khi chưa đạt. Giải quyết được việc "
        "lập trình viên và quản lý dự án không thể chung một ngưỡng.",
    )
    table(
        ["Nhóm", "Ngưỡng khỏe mạnh", "Nguồn"],
        [
            ["Nhân sự sản xuất trực tiếp", "70–90%", "Productive"],
            ["Quản lý khách hàng", "60–80%", "Productive"],
            ["Nhân sự mới vào nghề", "khoảng 75%", "Harvest"],
            ["Cấp quản lý", "thấp hơn trung bình là bình thường", "Productive"],
            ["Trung bình ngành dịch vụ", "65%", "Productive"],
        ],
        widths=[5.4, 6.0, 4.6],
    )
    note(
        "Productive cảnh báo: utilization liên tục cao là dấu hiệu quá sức và nguy "
        "cơ kiệt sức; utilization tụt đột ngột thường lộ vấn đề quản lý dự án hoặc "
        "giao tiếp.",
    )

    h("1.9 Hai thứ không sản phẩm nào làm", 2)
    field(
        "Số nhân sự theo thời gian",
        "không tìm thấy ở Productive, Float, Toggl, Harvest hay bất kỳ sản phẩm "
        "nào trong khảo sát. Đây là chỉ số nhân sự (HR), không phải chỉ số quản lý "
        "dịch vụ. Các sản phẩm này giả định số người đã biết, chỉ quan tâm người "
        "đó bận bao nhiêu. Nếu HeyE muốn, phải tự thiết kế.",
    )
    field(
        "Người mới bao lâu đạt năng suất",
        "cũng không sản phẩm nào làm. Thông lệ nhân sự đo bằng: số ngày từ ngày vào "
        "đến khi đạt ngưỡng năng suất và duy trì liên tục ít nhất hai tuần. Điều "
        "kiện “liên tục hai tuần” là thứ chống nhiễu — một tuần may mắn không tính.",
    )
