# -*- coding: utf-8 -*-
"""Phần 1: bìa, mục lục, bối cảnh, mô hình dữ liệu, khái niệm cốt lõi."""
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build(g):
    doc, h, p, rich, bullet, table, img, note, pagebreak = (
        g.doc, g.h, g.p, g.rich, g.bullet, g.table, g.img, g.note, g.pagebreak
    )
    BRAND, MUTED, INK, GOOD, BAD = g.BRAND, g.MUTED, g.INK, g.GOOD, g.BAD

    # ══════════════════ BÌA ══════════════════
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("HeyE — Kiểm soát chi phí")
    r.bold = True
    r.font.size = Pt(30)
    r.font.color.rgb = BRAND

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Tài liệu hướng dẫn tính năng và cấu hình")
    r.font.size = Pt(14)
    r.font.color.rgb = MUTED

    doc.add_paragraph()
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Cost Control Module — Feature & Configuration Guide")
    r.font.size = Pt(11)
    r.italic = True
    r.font.color.rgb = MUTED

    for _ in range(6):
        doc.add_paragraph()

    tb = doc.add_table(rows=5, cols=2)
    tb.alignment = 1
    meta = [
        ("Phiên bản", "v1.0"),
        ("Ngày phát hành", "24/08/2026"),
        ("Đối tượng đọc", "Toàn bộ team — PM, BA, Sale, Kế toán, Dev, QC"),
        ("Phạm vi", "Module Tài chính: 8 màn hình + cấu hình"),
        ("Trạng thái", "Đã hoàn thiện, đang chạy trên bản demo"),
    ]
    for i, (k, v) in enumerate(meta):
        c0, c1 = tb.rows[i].cells
        c0.text = ""
        rr = c0.paragraphs[0].add_run(k)
        rr.bold = True
        rr.font.size = Pt(10)
        c1.text = ""
        rr = c1.paragraphs[0].add_run(v)
        rr.font.size = Pt(10)
        c0.width, c1.width = Cm(4.5), Cm(11.5)

    pagebreak()

    # ══════════════════ MỤC LỤC ══════════════════
    h("Mục lục", 1)
    toc = [
        ("1", "Vấn đề đang giải quyết", "Vì sao cần module này"),
        ("2", "Mô hình dữ liệu — 5 khái niệm cốt lõi", "Hiểu đúng trước khi dùng"),
        ("3", "Hai trục tính tiền: Giá bán và Giá vốn", "Nền tảng của mọi con số"),
        ("4", "Màn 1 — Loại dịch vụ", "Danh mục gốc, khai một lần"),
        ("5", "Màn 2 — Bảng giá", "Đơn giá bán theo từng khách"),
        ("6", "Màn 3 — Giá vốn nhân sự", "Một giờ của mỗi người tốn bao nhiêu"),
        ("7", "Màn 4 — Hợp đồng", "Nơi ra lãi lỗ"),
        ("8", "Màn 5 — Giờ của tôi", "Nhân viên ghi giờ"),
        ("9", "Màn 6 — Giờ toàn công ty", "Quản lý duyệt và ghi hộ"),
        ("10", "Màn 7 — Chi phí", "Tiền chi ra ngoài lương"),
        ("11", "Màn 8 — Biểu đồ và báo cáo", "Đọc số ra quyết định"),
        ("12", "Quy tắc nghiệp vụ quan trọng", "Những điều dễ hiểu sai"),
        ("13", "Danh mục case dữ liệu mẫu", "Đối chiếu khi kiểm thử"),
        ("14", "Việc còn tồn và giới hạn", "Minh bạch phạm vi"),
    ]
    table(
        ["#", "Nội dung", "Tóm tắt"],
        [[a, f"**{b}**", c] for a, b, c in toc],
        widths=[1.2, 6.5, 8.7],
    )
    pagebreak()

    # ══════════════════ 1. VẤN ĐỀ ══════════════════
    h("1. Vấn đề đang giải quyết", 1)

    h("1.1 Trước khi có module này", 2)
    p(
        "Công ty bán phần mềm và dịch vụ triển khai cho ngân hàng. Mỗi hợp đồng "
        "kéo dài nhiều tháng, huy động nhiều người với mức lương khác nhau, kèm "
        "chi phí mua ngoài như bản quyền nền tảng, thuê thầu phụ, công tác phí."
    )
    p("Ba câu hỏi sau đây không ai trả lời được bằng số:")
    bullet("Hợp đồng Sacombank đang lãi hay lỗ, lãi bao nhiêu phần trăm?")
    bullet("Mảng Kiểm thử của cả công ty năm nay đóng góp bao nhiêu doanh thu?")
    bullet("Còn bao nhiêu ngân sách trước khi vượt hợp đồng?")
    p(
        "Không trả lời được vì ba mảnh dữ liệu nằm rời rạc: giá bán trong file "
        "báo giá, lương trong bảng lương, giờ làm trong đầu mỗi người. Ghép tay "
        "thì mất vài ngày, và ghép xong thì số đã cũ."
    )

    h("1.2 Sau khi có module này", 2)
    table(
        ["Câu hỏi", "Trả lời ở đâu", "Mất bao lâu"],
        [
            ["Hợp đồng này lãi lỗ ra sao", "Tab Tổng quan của hợp đồng", "Ngay lập tức"],
            ["Còn bao nhiêu ngân sách", "Ô Ngân sách, có thanh tiến độ", "Ngay lập tức"],
            ["Khi nào ngân sách cạn", "Cảnh báo đỏ trên biểu đồ", "Ngay lập tức"],
            ["Ai chưa ghi giờ tuần này", "Giờ toàn công ty, ô đỏ", "Ngay lập tức"],
            ["Chi phí nào chưa được duyệt", "Màn Chi phí, chip lọc", "Ngay lập tức"],
            ["Giờ nào làm mà không ra tiền", "Ô Chi phí, dòng Giờ không ra tiền", "Ngay lập tức"],
        ],
        widths=[6.2, 6.5, 3.7],
    )

    h("1.3 Nguyên tắc thiết kế", 2)
    note(
        "Module này dựng theo mô hình của Productive.io — phần mềm quản trị công ty "
        "dịch vụ được dùng rộng rãi trên thế giới. Cách sắp xếp màn hình, tên gọi "
        "và công thức tính đều bám theo chuẩn đó, nên mỗi nhãn tiếng Việt đều kèm "
        "thuật ngữ tiếng Anh gốc để đối chiếu khi cần.",
        "info",
    )
    pagebreak()

    # ══════════════════ 2. MÔ HÌNH DỮ LIỆU ══════════════════
    h("2. Mô hình dữ liệu — 5 khái niệm cốt lõi", 1)
    p(
        "Đọc kỹ phần này trước khi thao tác. Phần lớn lỗi nhập liệu đến từ việc "
        "nhầm lẫn giữa năm khái niệm dưới đây."
    )

    table(
        ["Khái niệm", "Là gì", "Ví dụ thật trong hệ thống", "Ai quản lý"],
        [
            [
                "**Loại dịch vụ**\nService type",
                "Danh mục LOẠI LAO ĐỘNG của công ty. Khai một lần, dùng cho mọi dự án.",
                "Thiết kế, Phát triển, Kiểm thử, Quản lý dự án, Triển khai, Phân tích yêu cầu, Bản quyền phần mềm",
                "Ban giám đốc",
            ],
            [
                "**Bảng giá**\nRate card",
                "ĐƠN GIÁ BÁN cho từng loại dịch vụ. Mỗi khách có thể có bảng riêng.",
                "Bảng giá Sacombank 2026 (7 dòng), Bảng giá OCB (6 dòng), Giá chuẩn công ty (6 dòng)",
                "Sale + Ban giám đốc",
            ],
            [
                "**Giá vốn**\nCost rate",
                "Một giờ của mỗi NGƯỜI tốn công ty bao nhiêu. Tính từ lương.",
                "Sơn: 39 triệu/tháng → 365.018 đ/giờ (đã gồm chi phí gián tiếp)",
                "Chỉ quản trị viên",
            ],
            [
                "**Hợp đồng**\nBudget",
                "Một lần bán cụ thể cho một khách. Chứa các HẠNG MỤC với số lượng và đơn giá.",
                "Phần mềm giao dịch STM — Giai đoạn 1 (Sacombank, 17 hạng mục, 4,88 tỷ)",
                "PM",
            ],
            [
                "**Hạng mục**\nService",
                "Một dòng bán trong hợp đồng. Đây là nơi DUY NHẤT ghi giờ và ghi chi phí vào được.",
                "Đọc chip CCCD qua NFC — 420 giờ × 750.000 đ = 315.000.000 đ",
                "PM",
            ],
        ],
        widths=[3.0, 4.3, 5.6, 3.5],
        small=True,
    )

    h("2.1 Quan hệ giữa chúng", 2)
    p(
        "Chuỗi phụ thuộc chạy từ trên xuống. Không có bước trước thì không làm được "
        "bước sau:"
    )
    note(
        "Loại dịch vụ  →  Bảng giá (gán đơn giá cho loại)  →  Hợp đồng  →  "
        "Hạng mục (rút giá từ bảng giá, thêm số lượng)  →  Ghi giờ / Ghi chi phí vào hạng mục",
        "info",
    )
    p(
        "Chiều ngược lại là chiều tính tiền: giờ ghi vào hạng mục sinh ra doanh thu "
        "theo đơn giá, đồng thời sinh ra chi phí theo giá vốn của người ghi. "
        "Chênh lệch giữa hai con số đó chính là lãi.",
        after=2,
    )

    h("2.2 Hai trục độc lập: Loại dịch vụ và Nhóm hạng mục", 2)
    p(
        "Đây là điểm hay bị nhầm nhất. Hai trục này vuông góc nhau, không thay thế "
        "được cho nhau:"
    )
    table(
        ["", "Loại dịch vụ", "Nhóm hạng mục (Section)"],
        [
            ["Trả lời câu hỏi", "LÀM VIỆC GÌ", "THUỘC PHẦN NÀO của sản phẩm"],
            ["Phạm vi", "Toàn công ty, dùng lại cho mọi hợp đồng", "Riêng từng hợp đồng"],
            [
                "Ví dụ",
                "Phát triển, Kiểm thử, Thiết kế",
                "Định danh & xác thực (eKYC), Giao dịch tiền mặt",
            ],
            ["Dùng để", "Báo cáo mảng nào lãi, ai làm loại việc gì", "Chia nhỏ hợp đồng cho dễ đọc"],
        ],
        widths=[3.6, 6.4, 6.4],
    )
    note(
        "Một hạng mục luôn có ĐỒNG THỜI cả hai: 'Đọc chip CCCD qua NFC' thuộc nhóm "
        "'Định danh & xác thực (eKYC)' và có loại dịch vụ là 'Phát triển'. "
        "Nhờ vậy hỏi được cả hai chiều: nhóm eKYC tốn bao nhiêu, và mảng Phát triển "
        "toàn công ty lãi bao nhiêu.",
        "good",
    )
    pagebreak()
