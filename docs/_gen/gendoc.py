# -*- coding: utf-8 -*-
"""
Sinh tài liệu DOCX hướng dẫn module Kiểm soát chi phí của HeyE.

Cấu trúc theo BA-PM-MasterSkill: BRD-style cho phần bối cảnh và mô hình,
PRD-style cho phần cấu hình từng màn, kèm bảng liệt kê case dữ liệu mock.
Mọi số liệu lấy từ DB thật, không bịa.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SHOTS = "C:/Users/dotha/AppData/Local/Temp/heye-doc"
OUT = "D:/hiendo25/heye-workflow/docs/HeyE-Kiem-soat-chi-phi-Huong-dan.docx"

BRAND = RGBColor(0x5B, 0x3D, 0xF5)
INK = RGBColor(0x1A, 0x1A, 0x2E)
MUTED = RGBColor(0x6B, 0x6B, 0x80)
GOOD = RGBColor(0x0E, 0x9F, 0x6E)
BAD = RGBColor(0xDC, 0x26, 0x26)

doc = Document()

# ── Thiết lập trang và font ─────────────────────────────────────────
for s in doc.sections:
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    s.left_margin = s.right_margin = Cm(2.0)
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(1.8)

st = doc.styles["Normal"]
st.font.name = "Segoe UI"
st.font.size = Pt(10.5)
st.font.color.rgb = INK
st._element.rPr.rFonts.set(qn("w:eastAsia"), "Segoe UI")
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.line_spacing = 1.15


def _style(name, size, bold=False, color=INK, before=14, after=6):
    s = doc.styles[name]
    s.font.name = "Segoe UI"
    s.font.size = Pt(size)
    s.font.bold = bold
    s.font.color.rgb = color
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True
    return s


_style("Heading 1", 18, True, BRAND, 22, 8)
_style("Heading 2", 14, True, INK, 16, 6)
_style("Heading 3", 11.5, True, INK, 12, 4)


def h(txt, lv=1):
    doc.add_heading(txt, lv)


def p(txt="", bold=False, italic=False, color=None, size=None, after=None):
    par = doc.add_paragraph()
    r = par.add_run(txt)
    r.bold, r.italic = bold, italic
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    if after is not None:
        par.paragraph_format.space_after = Pt(after)
    return par


def rich(parts, after=None):
    """parts = [(text, bold, color|None), ...]"""
    par = doc.add_paragraph()
    for t, b, c in parts:
        r = par.add_run(t)
        r.bold = b
        if c:
            r.font.color.rgb = c
    if after is not None:
        par.paragraph_format.space_after = Pt(after)
    return par


def bullet(txt, lv=0):
    par = doc.add_paragraph(txt, style="List Bullet")
    par.paragraph_format.left_indent = Cm(0.6 + lv * 0.6)
    par.paragraph_format.space_after = Pt(2)
    return par


def field(name, desc, required=False, num=None):
    """Một trường trong hộp thoại, viết theo kiểu trang help của Productive:
    gạch đầu dòng, tên trường in đậm, gạch ngang, rồi mô tả chạy liền.

    Bảng ba cột kín chữ đọc mệt hơn nhiều — mắt phải nhảy ngang liên tục.
    Dạng này đọc một mạch từ trên xuống.
    """
    # Có num thì dùng đoạn thường và tự đánh số, để số khớp với dấu trên ảnh.
    # Không có num thì dùng gạch đầu dòng.
    if num:
        par = doc.add_paragraph()
        par.paragraph_format.left_indent = Cm(0.6)
        par.paragraph_format.space_after = Pt(4)
        r = par.add_run(f"{num}. ")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = BRAND
    else:
        par = doc.add_paragraph(style="List Bullet")
        par.paragraph_format.left_indent = Cm(0.6)
        par.paragraph_format.space_after = Pt(4)
    r = par.add_run(name + (" ✱" if required else ""))
    r.bold = True
    r.font.size = Pt(10)
    r = par.add_run(" — " + desc)
    r.font.size = Pt(10)
    return par


def shade(cell, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def table(headers, rows, widths=None, small=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    size = Pt(8.5 if small else 9.5)
    for i, htxt in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(htxt)
        r.bold = True
        r.font.size = size
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        c.paragraphs[0].paragraph_format.space_after = Pt(1)
        shade(c, "5B3DF5")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            # Ô có thể nhiều dòng, mỗi dòng tự quyết định in đậm hay không.
            # Kiểm tra ** trên cả chuỗi thì dòng thứ hai làm hỏng phép so sánh,
            # kết quả là dấu ** lọt nguyên vào tài liệu.
            lines = str(v).splitlines() or [""]
            for k, line in enumerate(lines):
                par = cells[i].paragraphs[0] if k == 0 else cells[i].add_paragraph()
                par.paragraph_format.space_after = Pt(1)
                # Tách theo cặp ** để in đậm được cả cụm nằm GIỮA câu,
                # ví dụ "Doanh thu **+750.000 đ**" — không chỉ cả dòng.
                for j, seg in enumerate(line.split("**")):
                    if not seg:
                        continue
                    r = par.add_run(seg)
                    r.font.size = size
                    r.bold = j % 2 == 1  # đoạn lẻ nằm giữa cặp ** thì in đậm
                    if seg.startswith("✗") or seg.startswith("Thiếu"):
                        r.font.color.rgb = BAD
                    elif seg.startswith("✓"):
                        r.font.color.rgb = GOOD
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def img(fname, caption, width=16.4):
    path = os.path.join(SHOTS, fname)
    if not os.path.exists(path):
        p(f"[Thiếu ảnh: {fname}]", italic=True, color=BAD)
        return
    doc.add_picture(path, width=Cm(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED
    cap.paragraph_format.space_after = Pt(12)


def note(txt, kind="info"):
    """Khối ghi chú có nền màu."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.rows[0].cells[0]
    c.text = ""
    par = c.paragraphs[0]
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.space_before = Pt(3)
    r = par.add_run(txt)
    r.font.size = Pt(9.5)
    shade(c, {"info": "EEF0FE", "warn": "FEF6E7", "good": "EAF7F0"}[kind])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def pagebreak():
    doc.add_page_break()
